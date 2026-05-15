"""
DualRAG Services — Document Parser
====================================
Extracts raw text from PDF, DOCX, and TXT files.
Synchronous — called via ``asyncio.to_thread`` from async routes.
"""

from __future__ import annotations

import io
import logging
import re
from typing import Optional

logger = logging.getLogger("dualrag.services.parser")


def parse_document(file_bytes: bytes, extension: str) -> str:
    """
    Parse raw text from file bytes based on extension.

    Parameters
    ----------
    file_bytes : bytes
        The raw file content.
    extension : str
        Lowercased file extension including dot, e.g. ``".pdf"``.

    Returns
    -------
    str — Cleaned extracted text.

    Raises
    ------
    ValueError
        If the extension is unsupported or parsing fails.
    """
    ext = extension.lower()

    if ext == ".pdf":
        text = _parse_pdf(file_bytes)
    elif ext == ".docx":
        text = _parse_docx(file_bytes)
    elif ext == ".txt":
        text = _parse_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")

    cleaned = _clean_text(text)
    logger.info("Parsed %s document → %d chars", ext, len(cleaned))
    return cleaned


def _parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    import pdfplumber

    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
            else:
                logger.debug("Page %d yielded no text", i + 1)

    if not text_parts:
        raise ValueError("PDF contains no extractable text (possibly image-only)")

    return "\n\n".join(text_parts)


def _parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    text_parts: list[str] = []

    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            text_parts.append(stripped)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    if not text_parts:
        raise ValueError("DOCX contains no extractable text")

    return "\n\n".join(text_parts)


def _parse_txt(file_bytes: bytes) -> str:
    """Decode plain text, trying UTF-8 then Latin-1 fallback."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, ValueError):
            continue
    raise ValueError("Could not decode text file with any supported encoding")


def _clean_text(text: str) -> str:
    """Normalize whitespace and remove control characters."""
    # Replace tabs with spaces
    text = text.replace("\t", " ")
    # Collapse runs of 3+ newlines into 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Collapse runs of 3+ spaces into 1
    text = re.sub(r" {3,}", " ", text)
    # Strip leading/trailing whitespace
    text = text.strip()
    return text
